from faker import Faker
from datetime import datetime
import os
from docx import Document


fake = Faker("sl_SI")
os.makedirs("synthetic_docs", exist_ok=True)

# Preprosta predloga uradnega dopisa
template = """
REPUBLIKA SLOVENIJA
MINISTRSTVO ZA {ministrstvo}

Številka: {stevilka}
Datum: {datum}

Zadeva: {zadeva}

Spoštovani {prejemnik_ime},

Na podlagi {zakon} (Ur. l. RS, št. {ul_stevilka}) vas obveščamo o {vsebina}.

Lep pozdrav,

{posiljatelj_ime}
{posiljatelj_pozicija}
"""

for i in range(10):
    doc = template.format(
        ministrstvo=fake.job().split()[-1].upper(),
        stevilka=f"{fake.random_int(100,999)}/{datetime.now().year}",
        datum=datetime.now().strftime("%d.%m.%Y"),
        zadeva=fake.sentence(nb_words=6),
        prejemnik_ime=fake.name(),
        zakon="zakonu o varstvu osebnih podatkov",
        ul_stevilka=fake.random_int(1,100),
        vsebina=fake.paragraph(nb_sentences=3),
        posiljatelj_ime=fake.name(),
        posiljatelj_pozicija=f"Odgovorni {fake.job()}",
    )
    path = os.path.join("synthetic_docs", f"dopis_{i+1}.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(doc)
    print(f"Generirano: {path}")


fake = Faker("sl_SI")
os.makedirs("synthetic_docx", exist_ok=True)

for i in range(5):
    doc = Document()
    doc.add_heading('REPUBLIKA SLOVENIJA', level=1)
    doc.add_paragraph(f"MINISTRSTVO ZA {fake.job().split()[-1].upper()}")
    doc.add_paragraph(f"Številka: {fake.random_int(100,999)}/{datetime.now().year}")
    doc.add_paragraph(f"Datum: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_heading('Zadeva: ' + fake.sentence(nb_words=6), level=2)
    doc.add_paragraph("Spoštovani " + fake.name() + ",")
    doc.add_paragraph(fake.paragraph(nb_sentences=4))
    doc.add_paragraph("Lep pozdrav,")
    doc.add_paragraph(fake.name() + ", " + fake.job())

    path = os.path.join("synthetic_docx", f"dopis_{i+1}.docx")
    doc.save(path)
    print(f"Generirano: {path}")
