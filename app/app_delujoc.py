import streamlit as st
import json
import spacy
from io import BytesIO, StringIO
from pdfminer.high_level import extract_text
from docx import Document
from streamlit_js_eval import streamlit_js_eval
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
import re
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
import pandas as pd
from st_aggrid import AgGrid, GridOptionsBuilder, DataReturnMode, GridUpdateMode



# Registracija pisave z Unicode podporo
pdfmetrics.registerFont(TTFont("DejaVuSans", "D:\\faks\\fri-fu\\diploma\\aplikacija\\app\\DejaVuSans.ttf"))

# --- Setup ---
st.set_page_config(layout="wide")
nlp = spacy.load("D:\\faks\\fri-fu\\diploma\\aplikacija\\output\\model-last")
with open("D:\\faks\\fri-fu\\diploma\\aplikacija\\app\\config.json", encoding="utf-8") as f:
    label_map = json.load(f)

# --- Bralniki različnih formatov ---
def read_txt(buf): return buf.getvalue()
def read_docx(buf):
    doc = Document(buf)
    return "\n".join(p.text for p in doc.paragraphs)
def read_pdf(buf): return extract_text(buf)
def render_clickable_responsive1(text, selected_labels):
    doc = nlp(text)
    lm = {k: v for k, v in label_map.items() if k in selected_labels}

    # Zberi char-spans za vse entitete, ki jih želimo redaktirati
    char_spans = [(ent.start_char, ent.end_char) for ent in doc.ents if ent.label_ in lm]

    # Zgradimo HTML, kjer vsak znak ovijemo v <span>, in mu dodamo klaso 'redacted' če spada v kateri char-span
    html_chars = []
    for idx, ch in enumerate(text):
        cls = "token"
        # preveri, ali je ta znak znotraj katere izmed ent spanov
        if any(start <= idx < end for start, end in char_spans):
            cls += " redacted"
        # escapamo posebne znake
        esc = (ch.replace("&", "&amp;")
                  .replace("<", "&lt;")
                  .replace(">", "&gt;")
                  .replace(" ", "&nbsp;")
                  .replace("\n", "<br>"))
        html_chars.append(f'<span class="{cls}" data-idx="{idx}">{esc}</span>')

    html = f"""
    <style>
    #ctr {{
        background-color: white;
        padding: 1rem;
        border-radius: 6px;
        border: 1px solid #ccc;
        font-family: monospace;
        font-size: clamp(12px, 1.2vw, 18px);
        line-height: 1.6;
        max-height: 70vh;
        overflow-y: auto;
        overflow-x: hidden;
        white-space: normal;
        word-wrap: break-word;
        width: 100%;
        box-sizing: border-box;
    }}

    .token {{
        display: inline;
        cursor: pointer;
        word-break: break-word;
    }}

    .redacted {{
        background-color: black;
        color: transparent;
        border-radius: 2px;
    }}
    </style>
    <div id="ctr">{''.join(html_chars)}</div>
    <script>
    document.addEventListener("DOMContentLoaded", function() {{
        var container = document.getElementById('ctr');
        if (container) {{
            container.querySelectorAll('.token').forEach(function(span) {{
                span.addEventListener('click', function() {{
                    span.classList.toggle('redacted');
                }});
                span.style.height = "auto";
                span.style.height = span.scrollHeight + "px";
            }});
        }}
    }});
    </script>
    """
    st.components.v1.html(html, height=600, scrolling=True)

# --- Priprava HTML za interaktivno označevanje ---
def render_clickable_responsive(text, selected_labels):
    doc = nlp(text)
    lm = {k: v for k, v in label_map.items() if k in selected_labels}

    # Zberi char-spans za vse entitete, ki jih želimo redaktirati
    char_spans = [(ent.start_char, ent.end_char) for ent in doc.ents if ent.label_ in lm]

    # Pripravi seznam segmentov: vsak segment je tuple (start, end, is_redacted)
    segments = []
    idx = 0
    while idx < len(text):
        # ugotovi, če je trenutni znak redacted
        in_ent = any(s <= idx < e for s, e in char_spans)
        # poišči konec zaporedja z enako redacted-vrednostjo (in da ni presledek/nova vrstica)
        j = idx
        while j < len(text) and any(s <= j < e for s, e in char_spans) == in_ent and text[j] not in {" ", "\n"}:
            j += 1
        # zdaj segment [idx:j] ima istovrstno barvo
        segments.append((idx, j, in_ent))
        # posebej obravnaj presledke in newlines kot ločene segmente
        if j < len(text) and text[j] in {" ", "\n"}:
            segments.append((j, j+1, False))
            idx = j+1
        else:
            idx = j

    # Zgradi HTML
    html_parts = []
    for start, end, red in segments:
        chunk = text[start:end]
        # escapaj
        esc = (chunk.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                    .replace(" ", "&nbsp;")
                    .replace("\n", "<br>"))
        cls = "token"
        if red:
            cls += " redacted"
        html_parts.append(f'<span class="{cls}" data-idx="{start}">{esc}</span>')

    html = f"""
    <style>
    #ctr {{
        background-color: white;
        padding: 1rem;
        border-radius: 6px;
        border: 1px solid #ccc;
        font-family: monospace;
        font-size: clamp(12px, 1.2vw, 18px);
        line-height: 1.6;
        max-height: 70vh;
        overflow-y: auto;
        overflow-x: hidden;
        white-space: normal;
        word-wrap: break-word;
        width: 100%;
        box-sizing: border-box;
    }}

    .token {{
        display: inline;
        cursor: pointer;
        word-break: break-word;
    }}

    .redacted {{
        background-color: black;
        color: transparent;
        border-radius: 2px;
    }}
    </style>
    <div id="ctr">{''.join(html_parts)}</div>
    <script>
    document.addEventListener("DOMContentLoaded", function() {{
        var container = document.getElementById('ctr');
        if (container) {{
            container.querySelectorAll('.token').forEach(function(span) {{
                span.addEventListener('click', function() {{
                    span.classList.toggle('redacted');
                    // shrani zadnji klik v globalno spremenljivko
                    window.lastClickedSpan = {{
                    start: parseInt(span.dataset.start),
                    end: parseInt(span.dataset.end)
                    }};
                }});
                span.style.height = "auto";
                span.style.height = span.scrollHeight + "px";
            }});
        }}
    }});
    </script>
    """
    st.components.v1.html(html, height=600, scrolling=True)


# --- Ustvari PDF z ohranjenimi pravokotniki ---
def build_pdf(text, selected_labels, manual_idxs, font_name="DejaVuSans"):
    # 1) Najprej označimo bloke za redakcijo
    doc = nlp(text)
    lm = {k: v for k, v in label_map.items() if k in selected_labels}
    spans = set()
    # avtomatizirane entitete
    for ent in doc.ents:
        if ent.label_ in lm:
            spans.add((ent.start_char, ent.end_char))
    # ročno dodane
    spans |= set(tuple(mi) for mi in manual_idxs)

    # Sestavimo nov string, kjer so vsi spans zamenjani z blokom
    redacted = []
    idx = 0
    for start, end in sorted(spans):
        # besedilo pred spanom
        redacted.append(text[idx:start])
        length = max(end - start, 4)
        redacted.append("█" * length)
        idx = end
    redacted.append(text[idx:])
    redacted_text = "".join(redacted)

    # 2) Razbijmo na odstavke po dvojnih newline
    paragraphs = re.split(r"\n{2,}", redacted_text)

    # 3) Priprava ReportLab dokumenta
    buf = BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=30, rightMargin=30,
                            topMargin=30, bottomMargin=30)
    style = ParagraphStyle(
        name="Justified",
        fontName=font_name,
        fontSize=12,
        leading=14,
        alignment=TA_JUSTIFY,
        wordWrap='CJK'
    )

    story = []
    for para in paragraphs:
        # ohrani enojne prelome vrstice kot <br/>
        html_para = para.replace("\n", "<br/>")
        story.append(Paragraph(html_para, style))
        story.append(Spacer(1, 6))

    pdf.build(story)
    buf.seek(0)
    return buf.getvalue()

# --- DOCX builder ---
def build_docx(text, selected_labels, manual_idxs):
    doc_sp = nlp(text)
    lm = {k: v for k, v in label_map.items() if k in selected_labels}

    # merge entitites first
    with doc_sp.retokenize() as retok:
        for ent in doc_sp.ents:
            if ent.label_ in lm:
                retok.merge(ent, attrs={"LEMMA": lm[ent.label_], "ENT_TYPE": ent.label_})

    out = Document()
    para = out.add_paragraph()
    for tok in doc_sp:
        # newline
        if tok.text == "\n":
            para = out.add_paragraph()
            continue
        run = para.add_run(tok.text)
        if tok.ent_type_ in lm or tok.i in manual_idxs:
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), '000000')
            rPr.append(shd)
            run.font.color.rgb = RGBColor(0,0,0)
        # whitespace
        if tok.whitespace_:
            para.add_run(tok.whitespace_)
    buf = BytesIO()
    out.save(buf)
    return buf.getvalue()


# --- UI aplikacije ---
st.header("📄 Interaktivni anonimizator")
col1, col2 = st.columns([1, 3])

with col1:
    labels = list(label_map.keys())
    selected = st.multiselect("Izberite NER-entitete", labels, default=["PER", "LOC", "ORG"])
    uploaded = st.file_uploader("Naložite .txt/.docx/.pdf", type=["txt", "docx", "pdf"])
    if "manual" not in st.session_state:
        st.session_state.manual = []

    text = ""
    if uploaded and selected:
        ext = uploaded.name.split('.')[-1].lower()
        if ext == 'txt':
            text = read_txt(StringIO(uploaded.getvalue().decode('utf-8')))
        elif ext == 'docx':
            text = read_docx(uploaded)
        elif ext == 'pdf':
            text = read_pdf(uploaded)

        st.markdown("### 2. Prenos anonimizirane datoteke kot PDF")
        pdf_bytes = build_pdf(text, selected, st.session_state.manual)
        st.download_button(
            "⬇️ Prenesi PDF",
            data=pdf_bytes,
            file_name=f"{uploaded.name.rsplit('.',1)[0]}_anon.pdf",
            mime="application/pdf"
        )
    
    # # === NOV DEL: Izpišemo seznam entitet ===
    # doc = nlp(text)
    # lm = {k: v for k, v in label_map.items() if k in selected}
    # ents = [(ent.text, ent.label_) for ent in doc.ents if ent.label_ in lm]
    # if ents:
    #     st.markdown("### 2. Seznam označenih entitet")
    #     df = pd.DataFrame(ents, columns=["Entiteta", "Tip"])
    #     st.dataframe(df)  # prikaže tabelo z entitetami in njihovimi tipi
    # else:
    #     st.markdown("**Ni zaznanih entitet za izbrane tipe.**")

with col2:
    if uploaded and selected and text:
        st.markdown("### 1. Predogled besedila")
        st.markdown("> Kliknite na besedo, da jo (de)aktivirate kot ročno anonimizirano.")
        render_clickable_responsive(text, selected)
    
    # JavaScript, ki prebere window.lastClickedSpan in ga pošlje nazaj:
    js = """
    (() => {
        const span = window.lastClickedSpan;
        if (span) {
            // počistimo, da ne pošiljemo vedno istega
            window.lastClickedSpan = null;
            return JSON.stringify(span);
        }
        return null;
    })()
    """
    clicked = streamlit_js_eval.run_js(js)

    if clicked:
        span = json.loads(clicked)
        tup = (span["start"], span["end"])
        manual = st.session_state.manual
        # toggle: če že obstaja, odstrani, drugače dodaj
        if tup in manual:
            manual.remove(tup)
        else:
            manual.append(tup)
